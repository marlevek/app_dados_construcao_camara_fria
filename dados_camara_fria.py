import streamlit as st
from fpdf import FPDF
from docx import Document

st.set_page_config('Dados para Construção de Câmara Fria', page_icon=':material/warehouse:')
st.title(':blue[Levantamento de Dados para Construção de Câmara Fria] :material/warehouse:')

# Função para gerar PDF e Word
def gerar_pdf(dados):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font('Arial', size=12)
    pdf.cell(200, 10, txt="Levantamento de Dados para Construção de Câmara Fria", ln=True, align='C')
    pdf.ln(10)
    for key, value in dados.items():
        pdf.multi_cell(0, 10, f'{key}: {value}')
    return pdf.output(dest='S').encode('latin1')

def gerar_word(dados):
    doc = Document()
    doc.add_heading('Levantamento de Dados para Construção de Câmara Fria', 0)
    for key, value in dados.items():
        doc.add_paragraph(f'{key}: {value}')
    return doc

dados = {}
st.write('**Empresa Instalação: Climáts Ar-Condicionado e Refrigeração**')
dados['Empresa Instalação'] = 'Climáts Ar-Condicionado e Refrigeração'

dados['Cliente'] = st.text_input('**Cliente**')

painel = st.radio('**Tipo de Painel**', ['PIR', 'EPS'])
dados['Tipo de Painel'] = painel
if painel == 'PIR':
    dados['Espessura PIR'] = st.radio('PIR ºC', ['50mm 0ºC', '70mm -10ºC', '100mm -20ºC', '120mm -24ºC', '150mm -30ºC'])
else:
    dados['Espessura EPS'] = st.radio('EPS ºC', ['50mm +5ºC', '100mm -10ºC', '150mm -20ºC', '200mm -24ºC', '250mm -30ºC'])

dados['Isolamento do Piso'] = st.radio('**Isolamento do Piso:**', ['Painel', 'Civil Elevado', 'Civil Enterrado', 'Sem Isolamento'])
dados['Porta'] = st.radio('**Portas:**', ['Giratória 3 batentes', 'Giratória 4 batentes', 'Correr 3 batentes', 'Correr 4 batentes'])
dados['Quantidade de Portas'] = st.number_input('Quantidade de Portas:', min_value=1)
dados['Tamanho das Portas'] = st.radio('Tamanho das Portas:', ['Gir. 0,80 x 1,80', 'Gir. 1,00 x 2,00', 'Cor. 0,80 x 1,80', 'Cor. 1,00 x 2,00', 'Cor. 1,40 x 2,10', 'Cor. 1,40 x 2,40'])
dados['Lado Abertura da Porta'] = st.radio('Lado Abertura da Porta:', ['Direito', 'Esquerdo'])
dados['Produto Armazenado'] = st.selectbox('**Produto Armazenado**', ['Bebidas', 'Carne', 'Cerveja', 'Frutas', 'Gelo', 'Laticínio', 'Massa', 'Peixe', 'Sorvete', 'Vegetais', 'Outros'])

if dados['Produto Armazenado'] == 'Outros':
    dados['Especificação Produto'] = st.text_input('Especifique o tipo do produto: ')

dados['Temperatura de Entrada (ºC)'] = st.number_input('Temperatura de Entrada do Produto (ºC)', min_value=-30, value=0)
dados['Temperatura Interna Final (ºC)'] = st.number_input('Temperatura Interna Final (ºC)', min_value=-30, value=0)
dados['Tempo de Processo (h)'] = st.number_input('Tempo de Processo (h):', min_value=1)
dados['Capacidade Armazenada (kg)'] = st.number_input('Capacidade Armazenada (kg)', min_value=1)
dados['Movimentação'] = st.text_input('Movimentação')
dados['Distância Cond -> Evap (m)'] = st.number_input('**Distância entre unidade cond x evap (m)', min_value=1.0)
dados['Tensão Elétrica'] = st.radio('**Tensão Elétrica**', ['220V - monofásico', '220V - trifásico', '380 - trifásico'])

if st.button('Baixar PDF'):
    pdf_bytes = gerar_pdf(dados)
    st.download_button(label='Download PDF', data=pdf_bytes, file_name='dados_camara_fria.pdf', mime='application/pdf')
    
if st.button('Baixar Word'):
    doc = gerar_word(dados)
    doc_bytes = doc.save('dados_camara_fria.docx')
    with open('dados_camara_fria.docx', 'rb') as f:
        st.download_button(label='Download Word', data=f, file_name='dados_camara_fria.docx', mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

